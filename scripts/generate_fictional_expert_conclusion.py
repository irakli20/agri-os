from collections import Counter, defaultdict
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Inches


OUT = Path("/Users/macbookpro/საექსპერტო დასკვნები - აპრილი/საექსპ. დასკვნა 2605-121 საბურთალო - fictional training sample.docx")
LOGO = Path("/Users/macbookpro/agri-os/scripts/demaso_wood_1_logo.png")
EXPERT_SIGNATURE = Path("/Users/macbookpro/agri-os/scripts/original_signature_assets/drawing_1_image144.png")
DIRECTOR_SIGNATURE = Path("/Users/macbookpro/agri-os/scripts/original_signature_assets/drawing_5_image45.png")
COMPANY_STAMP = Path("/Users/macbookpro/agri-os/scripts/original_signature_assets/24d6489a2ad5_image73.png")
EXECUTION_COMPOSITE = Path("/Users/macbookpro/agri-os/scripts/original_signature_assets/execution_stamp_signature_composite.png")
FIRST_PAGE_SIGNATURE_BLOCK = Path("/Users/macbookpro/agri-os/scripts/original_signature_assets/first_page_signature_block.png")
GEORGIAN_FONT = Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf")


plants = [
    ("ლეღვი (Ficus carica)", "14-18", "3.5", "საღი/ ხილკენკრ."),
    ("ტყემალი (Prunus cerasifera)", "9.0", "3.0", "საღი/ ხილკენკრ."),
    ("ქლიავი (Prunus domestica)", "11.5", "3.2", "საღი/ ხილკენკრ."),
    ("ალუბალი (Prunus avium)", "13.0", "4.0", "საღი/ ხილკენკრ."),
    ("ბალი (Prunus avium)", "10.2", "3.4", "საღი/ ხილკენკრ."),
    ("ვაშლი (Malus domestica)", "16.0", "4.2", "საღი/ ხილკენკრ."),
    ("მსხალი (Pyrus communis)", "17.5", "4.5", "საღი/ ხილკენკრ."),
    ("ატამი (Prunus persica)", "8.5", "2.8", "საღი/ ხილკენკრ."),
    ("გარგარი (Prunus armeniaca)", "12.2", "3.6", "საღი/ ხილკენკრ."),
    ("ხურმა (Diospyros lotus)", "18.0", "5.0", "საღი/ ხილკენკრ."),
    ("კაკლის ხე (Juglans regia)", "22.0", "6.5", "საღი/ მაღალი ღირებულების ფოთლოვანი."),
    ("დაფნა (Laurus nobilis)", "7.5", "2.6", "საღი/ მაღალი ღირებულების ფოთლოვანი."),
    ("თუთა (Morus alba)", "19.0", "5.2", "საღი/ დაბალი ღირებულების ფოთლოვანი."),
    ("ცაცხვი (Tilia caucasica)", "21.0", "6.0", "საღი/ მაღალი ღირებულების ფოთლოვანი."),
    ("აკაცია (Robinia pseudoacacia)", "13.8", "4.3", "საღი/ დაბალი ღირებულების ფოთლოვანი."),
    ("აილანთო (Ailanthus altissima)", "12.4", "4.1", "საღი/ დაბალი ღირებულების ფოთლოვანი."),
    ("ნეკერჩხალი (Acer campestre)", "15.5", "5.0", "საღი/ მაღალი ღირებულების ფოთლოვანი."),
    ("იფანი (Fraxinus excelsior)", "18.5", "5.8", "საღი/ მაღალი ღირებულების ფოთლოვანი."),
    ("ჭადარი (Platanus orientalis)", "24.0", "7.0", "საღი/ მაღალი ღირებულების ფოთლოვანი."),
    ("ფიჭვი (Pinus eldarica)", "20.5", "6.5", "საღი/ წიწვოვანი."),
    ("კვიპაროსი (Cupressus sempervirens)", "16.0", "5.4", "საღი/ წიწვოვანი."),
    ("კედარი (Cedrus deodara)", "18.0", "5.8", "საღი/ წიწვოვანი."),
    ("ნაძვი (Picea orientalis)", "14.0", "4.6", "საღი/ წიწვოვანი."),
    ("ბზა (Buxus colchica)", "4.5", "1.6", "საღი/ ბუჩქი."),
    ("იასამანი (Syringa vulgaris)", "3-5", "2.0", "საღი/ ბუჩქი."),
    ("ვარდი (Rosa sp.)", "2-3", "1.2", "საღი/ ბუჩქი."),
    ("ბროწეული (Punica granatum)", "6.8", "2.7", "საღი/ ხილკენკრ."),
    ("ნუში (Prunus amygdalus)", "9.5", "3.1", "საღი/ ხილკენკრ."),
    ("კომში (Cydonia oblonga)", "10.5", "3.3", "საღი/ ხილკენკრ."),
    ("უნაბი (Zizyphus jujuba)", "8.0", "3.0", "საღი/ ხილკენკრ."),
    ("ლეღვი (Ficus carica)", "7-12", "2.8", "საღი/ ხილკენკრ."),
    ("ტყემალი (Prunus cerasifera)", "6.5", "2.4", "საღი/ ხილკენკრ."),
    ("ქლიავი (Prunus domestica)", "8.8", "2.9", "საღი/ ხილკენკრ."),
    ("ალუჩა (Prunus vachuschtii)", "7.2", "2.7", "საღი/ ხილკენკრ."),
    ("შინდი (Cornus mas)", "6.0", "2.5", "საღი/ ხილკენკრ."),
    ("თხილი (Corylus avellana)", "5-7", "2.4", "საღი/ ბუჩქი."),
    ("მუშმალა (Eriobotrya japonica)", "8.5", "2.8", "საღი/ ხილკენკრ."),
    ("ბალი (Prunus avium)", "9.0", "3.2", "საღი/ ხილკენკრ."),
    ("ვაშლი (Malus domestica)", "12.0", "3.7", "საღი/ ხილკენკრ."),
    ("მსხალი (Pyrus communis)", "11.0", "3.5", "საღი/ ხილკენკრ."),
    ("ატამი (Prunus persica)", "6.0", "2.3", "ხმობადი/ ხილკენკრ."),
    ("გარგარი (Prunus armeniaca)", "10.0", "3.0", "ზეხმელი/ ხილკენკრ."),
    ("აკაცია (Robinia pseudoacacia)", "14.0", "4.5", "ფაუტი/ დაბალი ღირებულების ფოთლოვანი."),
    ("აილანთო (Ailanthus altissima)", "9.0", "3.6", "ხმობადი/ დაბალი ღირებულების ფოთლოვანი."),
    ("ნაძვი (Picea orientalis)", "13.0", "4.4", "ზეხმელი/ წიწვოვანი."),
    ("ტუია (Thuja orientalis)", "8.5", "3.2", "საღი/ წიწვოვანი."),
    ("კვიპაროსი (Cupressus sempervirens)", "10.5", "4.0", "საღი/ წიწვოვანი."),
    ("ფიჭვი (Pinus eldarica)", "17.0", "5.6", "საღი/ წიწვოვანი."),
    ("ცაცხვი (Tilia caucasica)", "16.5", "5.2", "საღი/ მაღალი ღირებულების ფოთლოვანი."),
    ("იფანი (Fraxinus excelsior)", "15.0", "4.9", "საღი/ მაღალი ღირებულების ფოთლოვანი."),
    ("თუთა (Morus alba)", "12.5", "4.0", "საღი/ დაბალი ღირებულების ფოთლოვანი."),
    ("ალვის ხე (Populus nigra)", "20.0", "7.2", "საღი/ დაბალი ღირებულების ფოთლოვანი."),
    ("ბზა (Buxus colchica)", "3.8", "1.3", "საღი/ ბუჩქი."),
    ("იასამანი (Syringa vulgaris)", "4.2", "1.9", "საღი/ ბუჩქი."),
    ("ვარდი (Rosa sp.)", "2-3", "1.1", "საღი/ ბუჩქი."),
    ("ბროწეული (Punica granatum)", "5.5", "2.3", "საღი/ ხილკენკრ."),
    ("ნუში (Prunus amygdalus)", "7.8", "2.9", "საღი/ ხილკენკრ."),
    ("კომში (Cydonia oblonga)", "8.2", "3.0", "საღი/ ხილკენკრ."),
    ("უნაბი (Zizyphus jujuba)", "7.0", "2.8", "საღი/ ხილკენკრ."),
    ("ლეღვი (Ficus carica)", "9-13", "3.1", "საღი/ ხილკენკრ."),
]


meta = {
    "doc_no": "2604-811",
    "doc_date": "29.04.2026",
    "assessment_date": "24.04.2026",
    "assessment_date_text": "2026 წლის 24 აპრილს",
    "address": "თბილისი, საბურთალოს რაიონი, ვაჟა-ფშაველას გამზირი N45",
    "address_lines": "თბილისი, საბურთალოს რაიონი,\nვაჟა-ფშაველას გამზირი N45",
    "intro_location": "ქ. თბილისში, საბურთალოს რაიონში, ვაჟა-ფშაველას გამზირის N45 მისამართზე",
    "cadastre": "01.10.18.007.124",
    "owner": "შპს “საბურთალო დეველოპმენტი”",
    "owner_lines": "შპს “საბურთალო\nდეველოპმენტი”",
    "owner_id": "405321987",
    "area": "2140 მ²",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_no_wrap(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    no_wrap = tc_pr.first_child_found_in("w:noWrap")
    if no_wrap is None:
        no_wrap = OxmlElement("w:noWrap")
        tc_pr.append(no_wrap)


def set_table_grid(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells[: len(widths)]):
            set_cell_width(cell, widths[idx])


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    else:
        for child in list(borders):
            borders.remove(child)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "none")
        tag.set(qn("w:sz"), "0")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "FFFFFF")
        borders.append(tag)
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            cell_borders = tc_pr.first_child_found_in("w:tcBorders")
            if cell_borders is None:
                cell_borders = OxmlElement("w:tcBorders")
                tc_pr.append(cell_borders)
            else:
                for child in list(cell_borders):
                    cell_borders.remove(child)
            for edge in ("top", "left", "bottom", "right"):
                tag = OxmlElement(f"w:{edge}")
                tag.set(qn("w:val"), "none")
                tag.set(qn("w:sz"), "0")
                tag.set(qn("w:space"), "0")
                tag.set(qn("w:color"), "FFFFFF")
                cell_borders.append(tag)


def set_cell_text(cell, value, bold=False, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(value)
    run.bold = bold
    run.font.name = "Sylfaen"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Sylfaen")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0


def add_para(doc, txt="", bold=False, size=10, align=None, space_after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(txt)
    r.bold = bold
    r.font.name = "Sylfaen"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Sylfaen")
    r.font.size = Pt(size)
    return p


def add_paragraph_border(paragraph, top=False, bottom=False, size=6):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.first_child_found_in("w:pBdr")
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    for edge_name, enabled in (("top", top), ("bottom", bottom)):
        if not enabled:
            continue
        edge = OxmlElement(f"w:{edge_name}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), "4")
        edge.set(qn("w:color"), "000000")
        p_bdr.append(edge)


def add_horizontal_rule(doc, space_before=0, space_after=4, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    add_paragraph_border(p, bottom=True, size=size)
    return p


def add_image_paragraph(doc, image_path, width, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=0, left_indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    if left_indent is not None:
        p.paragraph_format.left_indent = left_indent
    if image_path.exists():
        run = p.add_run()
        run.add_picture(str(image_path), width=width)
    return p


def ensure_execution_composite():
    if EXECUTION_COMPOSITE.exists() or not (COMPANY_STAMP.exists() and DIRECTOR_SIGNATURE.exists()):
        return
    stamp = Image.open(COMPANY_STAMP).convert("RGBA")
    director = Image.open(DIRECTOR_SIGNATURE).convert("RGBA")
    canvas = Image.new("RGBA", (620, 310), (255, 255, 255, 0))
    director = director.resize((290, 290))
    stamp = stamp.resize((285, 285))
    canvas.alpha_composite(director, (35, 58))
    canvas.alpha_composite(stamp, (225, 0))
    canvas.save(EXECUTION_COMPOSITE)


def _font(size):
    from PIL import ImageFont

    if GEORGIAN_FONT.exists():
        return ImageFont.truetype(str(GEORGIAN_FONT), size=size)
    return ImageFont.load_default()


def _draw_text(draw, xy, text, font, fill=(0, 0, 0, 255), bold=False, italic=False):
    # PIL cannot reliably select SF Georgian bold/italic faces here, so use a subtle
    # stroke for the heavy labels and keep the block as an image for Google Docs.
    stroke = 1 if bold else 0
    draw.text(xy, text, font=font, fill=fill, stroke_width=stroke, stroke_fill=fill)


def ensure_first_page_signature_block():
    if FIRST_PAGE_SIGNATURE_BLOCK.exists():
        return
    ensure_execution_composite()
    from PIL import ImageDraw

    canvas = Image.new("RGBA", (1180, 285), (255, 255, 255, 0))
    draw = ImageDraw.Draw(canvas)
    label_font = _font(28)
    text_font = _font(27)
    small_font = _font(24)

    _draw_text(draw, (5, 18), "ექსპერტი:", label_font, bold=True)
    if EXPERT_SIGNATURE.exists():
        sig = Image.open(EXPERT_SIGNATURE).convert("RGBA").resize((210, 96))
        canvas.alpha_composite(sig, (165, 0))
    _draw_text(draw, (395, 20), "ლევან გოშაძე (01008013049)", text_font, bold=True)

    left_lines = [
        ("შემსრულებელი:", True, 95),
        ("შპს დემასო ვუდ 1 (400331716)", True, 133),
        ("დირექტორი: ემზარ შარია", True, 171),
        ("info@demaso.ge", False, 209),
        ("(+995) 593 18 18 70", False, 247),
    ]
    for txt, bold, y in left_lines:
        _draw_text(draw, (5, y), txt, small_font, bold=bold)

    if DIRECTOR_SIGNATURE.exists():
        director = Image.open(DIRECTOR_SIGNATURE).convert("RGBA").resize((245, 176))
        canvas.alpha_composite(director, (370, 102))
    if COMPANY_STAMP.exists():
        stamp = Image.open(COMPANY_STAMP).convert("RGBA").resize((215, 215))
        canvas.alpha_composite(stamp, (530, 66))

    canvas.save(FIRST_PAGE_SIGNATURE_BLOCK)


def add_first_page_signature_block(doc):
    ensure_first_page_signature_block()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if FIRST_PAGE_SIGNATURE_BLOCK.exists():
        p.add_run().add_picture(str(FIRST_PAGE_SIGNATURE_BLOCK), width=Inches(6.5))


def heading(doc, txt):
    p = add_para(doc, txt, bold=True, size=13, space_after=6)
    p.paragraph_format.space_before = Pt(12)
    return p


def add_contents_block(doc):
    rows = [
        ("შესავალი", "1"),
        ("ტერიტორიის აღწერა", "1"),
        ("ჩატარებული სამუშაო", "1"),
        ("რაოდენობები და ნომრები სახეობების მიხედვით", "1"),
        ("მცენარეთა მდგომარეობა", "3"),
        ("დასკვნა", "3"),
    ]
    spacer = add_para(doc, "", size=10, space_after=0)
    spacer.paragraph_format.line_spacing = 1.0
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.autofit = False
    set_table_grid(table, [6200, 900])
    remove_table_borders(table)
    set_cell_text(table.cell(0, 0), "დოკუმენტის შემადგენლობა", bold=True, size=13)
    table.cell(0, 1).text = ""
    for idx, (label, page) in enumerate(rows, start=1):
        for col, value in enumerate((label, page)):
            cell = table.cell(idx, col)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if col == 1 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            run.font.name = "Sylfaen"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Sylfaen")
            run.font.size = Pt(12)
    add_para(doc, "", size=10, space_after=0)
    add_para(doc, "დანართი 1. მერქნიან მცენარეთა შეფასების უწყისი", bold=True, size=11, space_after=0)
    link = add_para(doc, "ფოტოდოკუმენტაცია:", bold=True, size=11, space_after=0)
    url = add_para(doc, "https://1drv.ms/f/c/e8c36b8b2ac24e4c/IgCLGXSDteX5S5F1K5G40HCZAW3EWPiYey28gD2eTOwkzg0", size=10, space_after=6)
    for run in url.runs:
        run.underline = True
    add_para(doc, "", size=10, space_after=0)


def add_property_summary(doc):
    table = doc.add_table(rows=2, cols=4)
    table.autofit = False
    set_table_grid(table, [4000, 2300, 2200, 1500])
    remove_table_borders(table)
    labels = ["ტერიტორიის მისამართი", "ს/კ", "მესაკუთრე", "ფართობი"]
    values = [meta.get("address_lines", meta["address"]), meta["cadastre"], meta.get("owner_lines", meta["owner"]), meta["area"]]
    for idx, label in enumerate(labels):
        cell = table.cell(0, idx)
        if idx in (1, 3):
            set_cell_no_wrap(cell)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if idx == 3 else WD_ALIGN_PARAGRAPH.LEFT
        if idx == 3:
            p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(0)
        if idx == 2:
            p.paragraph_format.left_indent = Cm(0.25)
        run = p.add_run(label)
        run.bold = True
        run.italic = True
        run.font.name = "Sylfaen"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Sylfaen")
        run.font.size = Pt(10.5 if idx == 3 else 12)
        add_paragraph_border(p, bottom=True, size=8)
    for idx, value in enumerate(values):
        cell = table.cell(1, idx)
        if idx in (1, 3):
            set_cell_no_wrap(cell)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if idx == 3 else WD_ALIGN_PARAGRAPH.LEFT
        if idx == 3:
            p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(0)
        if idx == 2:
            p.paragraph_format.left_indent = Cm(0.25)
        run = p.add_run(value)
        run.italic = True
        run.font.name = "Sylfaen"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Sylfaen")
        run.font.size = Pt(10.5 if idx == 3 else 11)
    add_horizontal_rule(doc, space_before=2, space_after=4, size=12)


def add_annex_metadata(doc):
    rows = [
        ("მისამართი:", meta["address"]),
        ("საკადასტრო კოდი:", meta["cadastre"]),
        ("ფართობი:", meta["area"]),
        ("მესაკუთრე:", f"{meta['owner']} (ID № {meta['owner_id']})"),
        ("შეფასების მეთოდი:", "ვიზუალური"),
        ("შეფასების თარიღი:", meta["assessment_date"]),
        ("შემსრულებელი:", "ლ. გოშაძე, ე. შარია"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    set_table_grid(table, [2175, 7905])
    remove_table_borders(table)
    for r, (label, value) in enumerate(rows):
        label_cell = table.cell(r, 0)
        value_cell = table.cell(r, 1)
        for cell, txt, bold in ((label_cell, label, True), (value_cell, value, False)):
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(txt)
            run.bold = bold
            run.italic = True if bold else False
            run.font.name = "Sylfaen"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Sylfaen")
            run.font.size = Pt(10)
            add_paragraph_border(p, bottom=True, size=8)


def category(desc):
    if "ხილკენკრ" in desc:
        return "ხილკენკროვანი"
    if "ბუჩქ" in desc:
        return "ბუჩქი"
    if "წიწვოვანი" in desc:
        return "წიწვოვანი"
    if "მაღალი ღირებულების" in desc:
        return "მაღალი ღირებულების ფოთლოვანი"
    if "დაბალი ღირებულების" in desc:
        return "დაბალი ღირებულების ფოთლოვანი"
    return "სხვა"


def condition(desc):
    for key in ["ზეხმელი", "ხმობადი", "ფაუტი", "ამორტიზებული"]:
        if key in desc:
            return key
    return "ჯანსაღი"


def species_name(full):
    return full.split(" (")[0]


def condition_counts_text(plants_data):
    counts = Counter(condition(p[3]) for p in plants_data)
    healthy = counts.pop("ჯანსაღი", 0)
    parts = [f"{healthy} არის ჯანსაღი"]
    for key in ["ზეხმელი", "ხმობადი", "ფაუტი", "ამორტიზებული"]:
        if counts.get(key):
            parts.append(f"{counts[key]} - {key}")
    for key, value in counts.items():
        if value:
            parts.append(f"{value} - {key}")
    return ", ".join(parts)


def unhealthy_parenthetical(plants_data):
    counts = Counter(condition(p[3]) for p in plants_data if condition(p[3]) != "ჯანსაღი")
    ordered = []
    for key in ["ზეხმელი", "ხმობადი", "ფაუტი", "ამორტიზებული"]:
        if counts.get(key):
            ordered.append(f"{counts[key]} {key}")
    for key, value in counts.items():
        if key not in {"ზეხმელი", "ხმობადი", "ფაუტი", "ამორტიზებული"}:
            ordered.append(f"{value} {key}")
    return ", ".join(ordered)


def add_table(doc, rows, cols, widths=None, header=True):
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    table.autofit = False
    for row in table.rows:
        for cell in row.cells:
            cell.margin_top = 80
            cell.margin_bottom = 80
            cell.margin_left = 80
            cell.margin_right = 80
    if widths:
        set_table_grid(table, widths)
    if header:
        for cell in table.rows[0].cells:
            set_cell_shading(cell, "D9EAF7")
    return table


def diameter_13m(index, diameter_10cm):
    if "-" in diameter_10cm:
        return ""
    try:
        value = float(diameter_10cm)
    except ValueError:
        return ""
    if index % 5 == 0:
        return ""
    adjusted = max(1.5, value - (1.8 if value > 10 else 1.0))
    return f"{adjusted:.1f}".rstrip("0").rstrip(".")


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.3)
    sec.bottom_margin = Cm(1.3)
    sec.left_margin = Cm(1.5)
    sec.right_margin = Cm(1.5)

    styles = doc.styles
    styles["Normal"].font.name = "Sylfaen"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Sylfaen")
    styles["Normal"].font.size = Pt(10)

    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    set_table_grid(header_table, [2500, 7700])
    remove_table_borders(header_table)
    left = header_table.cell(0, 0)
    left.text = ""
    left_p = left.paragraphs[0]
    left_p.paragraph_format.space_before = Pt(12)
    left_p.paragraph_format.space_after = Pt(0)
    run = left_p.add_run(f"დოკ. №: {meta['doc_no']}")
    run.bold = True
    run.font.name = "Sylfaen"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Sylfaen")
    run.font.size = Pt(9)
    left_p.add_run("\n")
    date_run = left_p.add_run(f"თარიღი: {meta['doc_date']}")
    date_run.font.name = "Sylfaen"
    date_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Sylfaen")
    date_run.font.size = Pt(9)
    right = header_table.cell(0, 1)
    right.text = ""
    logo_p = right.paragraphs[0]
    logo_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    logo_p.paragraph_format.space_before = Pt(0)
    logo_p.paragraph_format.space_after = Pt(0)
    logo_p.paragraph_format.line_spacing = 1.0
    if LOGO.exists():
        logo_p.add_run().add_picture(str(LOGO), width=Cm(3.0))
    add_horizontal_rule(doc, space_before=0, space_after=2, size=16)
    title = add_para(doc, "მწვანე ნარგავების ხარისხობრივი შეფასება და\nსაექსპერტო დასკვნა", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)
    for run in title.runs:
        run.italic = True
    title.paragraph_format.space_before = Pt(4)
    add_paragraph_border(title, bottom=True, size=12)

    add_property_summary(doc)

    plant_count = len(plants)
    unhealthy_text = unhealthy_parenthetical(plants)
    condition_text = condition_counts_text(plants)

    add_para(doc, f"შეფასებული მცენარეები: {plant_count}", bold=True, size=11)

    add_contents_block(doc)

    heading(doc, "შესავალი")
    add_para(doc, f"დოკუმენტში ასახულია {meta['intro_location']} რეგისტრირებულ, ს/კ {meta['cadastre']} მიწის ნაკვეთზე არსებულ მერქნიან მცენარეთა ხარისხობრივი შეფასების შედეგები.")
    add_para(doc, "იგი მომზადებულია ქ. თბილისის მუნიციპალიტეტის საკრებულოს 11.06.2022 №11-58 დადგენილებისა და საქართველოს 01.12.2000 №594 კანონის მოთხოვნათა შესაბამისად.")
    add_para(doc, "დოკუმენტს დანართებად ახლავს მცენარეთა შეფასების უწყისი და ფოტოსურათები.")

    heading(doc, "ტერიტორიის აღწერა")
    add_para(doc, f"მიწის ნაკვეთს გააჩნია არასასოფლო-სამეურნეო დანიშნულება და წარმოადგენს კერძო საკუთრებას. მისი ფართობია {meta['area']}.")
    add_para(doc, "თბილისის ფუნქციური ზონირების მიხედვით, ნაკვეთი მდებარეობს შერეულ საცხოვრებელ ზონაში და მასზე წარმოდგენილია სხვადასხვა ასაკისა და მდგომარეობის მერქნიანი მცენარეები.")

    heading(doc, "ჩატარებული სამუშაო")
    add_para(doc, f"{meta['assessment_date_text']}, ვიზუალური დათვალიერების გზით ჩატარდა ს/კ {meta['cadastre']} მიწის ნაკვეთზე არსებული მერქნიანი მცენარეების შეფასება.")
    add_para(doc, "შეფასებისას, მცენარეთა აღწერა მოხდა ინდივიდუალურად (განცალკევებით).")
    add_para(doc, "შეფასებისას, განცალკევებით აღწერილ თითოეულ ინდივიდს მიენიჭა უნიკალური ნომერი. თითოეული მცენარე ამ ნომრით მოიხსენიება წინამდებარე დოკუმენტში და მიწის ნაკვეთის ტოპო-გეგმაზე.")
    add_para(doc, "თითოეული ერთეულისთვის აღებულ იქნა სატაქსაციო მონაცემები, აღიწერა მდგომარეობა, შეფასდა გადარგვა-გახარების შესაძლებლობა, მოხდა თითოეულის დასურათება.")
    add_para(doc, "მცენარეთა აღწერის საბაზისო მონაცემები ასახულია მცენარეთა შეფასების უწყისში (იხ. დანართი 1).")
    add_para(doc, "წინამდებარე დოკუმენტში, „მცენარის დიამეტრში“ იგულისხმება მიწიდან დაახლოებით 10 სმ სიმაღლეზე გაზომილი მცენარის ღეროს დიამეტრი.")

    counts_by_cat = Counter(category(p[3]) for p in plants)
    unhealthy = [(i + 1, p) for i, p in enumerate(plants) if condition(p[3]) != "ჯანსაღი"]
    red_list = [(11, plants[10]), (12, plants[11]), (24, plants[23])]
    transplantable = [1, 2, 3, 4, 5, 6, 8, 9, 27, 28, 29, 30, 31, 32, 33, 34, 35, 36, 37, 38, 39, 40, 46, 47, 48, 56, 57, 58, 59, 60]

    heading(doc, "რაოდენობები და ნომრები სახეობების მიხედვით")
    add_para(doc, f"შეფასებული {plant_count} მერქნიანი მცენარე სახეობებისა და ნომრების მიხედვით შემდეგნაირად ნაწილდება:")
    for cat, count in counts_by_cat.items():
        nums = [str(i + 1) for i, p in enumerate(plants) if category(p[3]) == cat]
        add_para(doc, f"{count} {cat} - NN {', '.join(nums)}")

    heading(doc, "მცენარეთა მდგომარეობა")
    add_para(doc, f"შეფასების შედეგად დადგინდა, რომ {plant_count} მერქნიანი მცენარიდან {condition_text}.")
    bad_table = add_table(doc, len(unhealthy) + 1, 4)
    for i, h in enumerate(["N", "სახეობა", "მდგომარეობა", "მცენარის ნომერი"]):
        set_cell_text(bad_table.cell(0, i), h, bold=True, size=9)
    for r, (num, p) in enumerate(unhealthy, start=1):
        set_cell_text(bad_table.cell(r, 0), str(r), size=9)
        set_cell_text(bad_table.cell(r, 1), species_name(p[0]), size=9)
        set_cell_text(bad_table.cell(r, 2), condition(p[3]), size=9)
        set_cell_text(bad_table.cell(r, 3), f"N {num}", size=9)

    heading(doc, "გადარგვის შესაძლებლობა")
    add_para(doc, "გადარგვის შესაძლებლობის განსაზღვრისას მხედველობაში მიიღება შემდეგი პირობები: მცენარე უნდა იყოს ჯანსაღი და კარგად განვითარებული; არ იყოს წვერხმელი; არ აღენიშნებოდეს გულის სიდამპლე ან სხვა დაავადება; ვარჯში ხმელი ტოტების რაოდენობა არ აღემატებოდეს 15 %-ს და არ იყოს გადაბერებული.")
    add_para(doc, f"შეფასების შედეგად დადგინდა, რომ, ზემოთ აღნიშნული კრიტერიუმებით, საჭიროების შემთხვევაში, სათანადო მეთოდიკის გამოყენებით, გადარგვას შესაძლებელია დაექვემდებაროს 30 შეფასებული მცენარე - NN {', '.join(map(str, transplantable))}.")

    heading(doc, "დასკვნა")
    add_para(doc, f"ქვემოთ, პუნქტებად მოცემულია {meta['assessment_date_text'].replace('წლის', 'წლის')}, {meta['intro_location']} მდებარე მიწის ნაკვეთზე, {plant_count} მერქნიანი მცენარის ხარისხობრივი შეფასების შემაჯამებელი მონაცემები.")
    add_para(doc, f"1. რაოდენობები სახეობათა ჯგუფებისა და მცენარეთა მდგომარეობის მიხედვით. შეფასებული {plant_count} მერქნიანი მცენარიდან {condition_text}. მათ შორის არის:")
    for cat, count in counts_by_cat.items():
        add_para(doc, f"{count} {cat};")
    add_para(doc, "2. საქართველოს წითელი ნუსხით დაცულ მცენარეთა სახეობები, რაოდენობები, ნომრები. აღიწერა საქართველოს წითელი ნუსხით დაცული 3 მცენარე, კერძოდ:")
    add_para(doc, "1 კაკლის ხე (Juglans regia) N 11")
    add_para(doc, "1 დაფნა (Laurus nobilis) N 12")
    add_para(doc, "1 ბზა (Buxus colchica) N 24")
    add_para(doc, f"3. არაჯანსაღ მცენარეთა სახეობები, რაოდენობები, ნომრები. სულ გამოვლინდა {len(unhealthy)} არაჯანსაღი ({unhealthy_text}) მცენარე, კერძოდ:")
    for num, p in unhealthy:
        add_para(doc, f"1 {species_name(p[0])} N {num} ({condition(p[3])})")
    add_para(doc, "4. გადარგვადი მცენარეები. გადარგვის შესაძლებლობის განსაზღვრის კრიტერიუმების მიხედვით, სათანადო მეთოდიკის გამოყენებით, გადარგვას შესაძლებელია დაექვემდებაროს 30 შეფასებული მცენარე.")
    add_para(doc, "მეტი დეტალისთვის იხილეთ:")
    add_para(doc, "დანართი 1. მერქნიან მცენარეთა შეფასების უწყისი;")
    add_para(doc, "დანართი 2. ფოტოდოკუმენტაცია;")
    add_para(doc, "დამკვეთის მიერ თბილისის მიერიაში ამ დასკვნასთან ერთად წარდგენილი ტოპო-გეგმა.")

    add_para(doc, "")
    add_first_page_signature_block(doc)

    doc.add_page_break()
    heading(doc, f"დანართი 1. მერქნიან მცენარეთა შეფასების უწყისი - ნაკვ. {meta['cadastre']}")
    add_annex_metadata(doc)

    ledger_widths = [435, 3120, 810, 750, 660, 4305]
    ledger = add_table(doc, len(plants) + 3, 6, widths=ledger_widths)
    headers = ["№", "მცენარის დასახელება", "დიამეტრი (სმ) მიწიდან:", "", "სიმაღლე (მ)", "აღწერა"]
    for c, h in enumerate(headers):
        set_cell_text(ledger.cell(0, c), h, bold=True, size=7.5)
    ledger.cell(0, 2).merge(ledger.cell(0, 3))
    set_cell_text(ledger.cell(1, 0), "", size=7.5)
    set_cell_text(ledger.cell(1, 1), "", size=7.5)
    set_cell_text(ledger.cell(1, 2), "10 სმ", bold=True, size=7.5)
    set_cell_text(ledger.cell(1, 3), "1.3 მ", bold=True, size=7.5)
    set_cell_text(ledger.cell(1, 4), "", size=7.5)
    set_cell_text(ledger.cell(1, 5), "", size=7.5)
    for cell in ledger.rows[1].cells:
        set_cell_shading(cell, "D9EAF7")
    set_repeat_table_header(ledger.rows[0])
    set_repeat_table_header(ledger.rows[1])
    for r, p in enumerate(plants, start=1):
        row_idx = r + 1
        set_cell_text(ledger.cell(row_idx, 0), str(r), size=7.5)
        set_cell_text(ledger.cell(row_idx, 1), p[0], size=7.5)
        set_cell_text(ledger.cell(row_idx, 2), p[1], size=7.5)
        set_cell_text(ledger.cell(row_idx, 3), diameter_13m(r, p[1]), size=7.5)
        set_cell_text(ledger.cell(row_idx, 4), p[2], size=7.5)
        set_cell_text(ledger.cell(row_idx, 5), p[3], size=7.5)
    note_row = len(plants) + 2
    set_cell_text(ledger.cell(note_row, 0), "*", size=7.5)
    merged = ledger.cell(note_row, 1).merge(ledger.cell(note_row, 5))
    set_cell_text(merged, "გადარგვადად მიიჩნევა მცენარე, რომელიც არის ჯანსაღი და კარგად განვითარებული, არ არის წვერხმელი, არ აღენიშნება გულის სიდამპლე ან სხვა დაავადება და არ არის გადაბერებული/ამორტიზებული.", size=7.5)

    add_para(doc, "")
    sig2 = doc.add_table(rows=1, cols=6)
    sig2.autofit = False
    set_table_grid(sig2, [1500, 1350, 1250, 1500, 1700, 1800])
    remove_table_borders(sig2)
    final_cells = [
        ("ლევან გოშაძე", None),
        ("", EXPERT_SIGNATURE),
        ("ექსპერტი;", None),
        ("ემზარ შარია", None),
        ("", DIRECTOR_SIGNATURE),
        ("დირექტორი,\nშპს დემასო ვუდ 1", None),
    ]
    for idx, (txt, image) in enumerate(final_cells):
        cell = sig2.cell(0, idx)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if image and image.exists():
            width = Inches(1.45) if image == EXPERT_SIGNATURE else Inches(1.7)
            height = Inches(0.62) if image == EXPERT_SIGNATURE else Inches(0.65)
            p.add_run().add_picture(str(image), width=width, height=height)
        else:
            run = p.add_run(txt)
            run.bold = True
            run.font.name = "Sylfaen"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Sylfaen")
            run.font.size = Pt(10)

    doc.add_page_break()
    heading(doc, f"ფოტოდოკუმენტაცია - ნაკვ. {meta['cadastre']}")
    add_para(doc, "ფოტოდოკუმენტაცია წარმოდგენილი იქნება ცალკე დანართის სახით. ეს სასწავლო ნიმუში შექმნილია სურათების გარეშე, ტექსტური სტრუქტურისა და შეფასების უწყისის მოდელირების მიზნით.", size=10)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
